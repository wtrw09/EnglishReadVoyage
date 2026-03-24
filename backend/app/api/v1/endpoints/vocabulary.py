"""生词本API端点。"""
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc, and_, delete
from typing import List
import io
import os

from app.core.database import get_db
from app.api.dependencies import get_current_user
from app.models.database_models import User, Vocabulary
from app.schemas.vocabulary import VocabularyCreate, VocabularyResponse, VocabularyListResponse, VocabularyBatchDelete, VocabularyExport
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 获取字体文件路径
FONT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'fonts', 'simsun.ttc')
TITLE_FONT_NAME = '方正小标宋简体'  # 标题字体
TABLE_FONT_NAME = '黑体'  # 表格字体


def set_run_font(run, font_name, font_size=None):
    """设置run的字体"""
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    if font_size:
        run.font.size = font_size


def set_cell_vertical_center(cell):
    """设置表格单元格垂直居中"""
    # 使用python-docx的内置方法，需要使用枚举值
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


router = APIRouter()


@router.post("/", response_model=VocabularyResponse)
async def add_vocabulary(
    request: VocabularyCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    添加生词到生词本。

    如果相同的单词和句子组合已存在，则返回已存在的记录。
    """
    # 检查是否已存在相同的生词（同一用户、同一单词、同一句子）
    stmt = select(Vocabulary).where(
        and_(
            Vocabulary.user_id == current_user.id,
            Vocabulary.word == request.word,
            Vocabulary.sentence == request.sentence
        )
    )
    result = await db.execute(stmt)
    existing = result.scalar_one_or_none()

    if existing:
        # 已存在则返回已存在的记录
        return existing

    # 创建新的生词记录
    vocab = Vocabulary(
        user_id=current_user.id,
        word=request.word,
        phonetic=request.phonetic,
        translation=request.translation,
        sentence=request.sentence,
        book_name=request.book_name
    )
    db.add(vocab)
    await db.commit()
    await db.refresh(vocab)

    return vocab


@router.get("/", response_model=VocabularyListResponse)
async def get_vocabulary_list(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    获取当前用户的生词列表。

    按添加时间倒序排列。
    """
    stmt = select(Vocabulary).where(
        Vocabulary.user_id == current_user.id
    ).order_by(desc(Vocabulary.created_at))

    result = await db.execute(stmt)
    vocab_list = result.scalars().all()

    return VocabularyListResponse(
        items=[VocabularyResponse.model_validate(v) for v in vocab_list],
        total=len(vocab_list)
    )


@router.delete("/{vocab_id}")
async def delete_vocabulary(
    vocab_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    删除生词本中的生词。

    只能删除自己添加的生词。
    """
    stmt = select(Vocabulary).where(
        and_(
            Vocabulary.id == vocab_id,
            Vocabulary.user_id == current_user.id
        )
    )
    result = await db.execute(stmt)
    vocab = result.scalar_one_or_none()

    if not vocab:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="生词不存在"
        )

    await db.delete(vocab)
    await db.commit()

    return {"success": True, "message": "生词已删除"}


@router.get("/check", response_model=dict)
async def check_vocabulary_exists(
    word: str,
    sentence: str = "",
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    检查生词是否已存在于生词本中。
    """
    stmt = select(Vocabulary).where(
        and_(
            Vocabulary.user_id == current_user.id,
            Vocabulary.word == word,
            Vocabulary.sentence == sentence
        )
    )
    result = await db.execute(stmt)
    existing = result.scalar_one_or_none()

    return {
        "exists": existing is not None,
        "id": existing.id if existing else None
    }


@router.post("/batch-delete")
async def batch_delete_vocabulary(
    request: VocabularyBatchDelete,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    批量删除生词本中的生词。

    只能删除自己添加的生词。
    """
    ids = request.ids
    if not ids:
        return {"success": True, "deleted_count": 0}

    # 删除属于当前用户的指定生词
    stmt = delete(Vocabulary).where(
        and_(
            Vocabulary.id.in_(ids),
            Vocabulary.user_id == current_user.id
        )
    )
    result = await db.execute(stmt)
    await db.commit()

    return {"success": True, "deleted_count": result.rowcount}


@router.post("/export")
async def export_vocabulary(
    request: VocabularyExport,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    导出生词本为Word文档。

    可以指定要隐藏的字段（用于默写练习）：
    - word: 英文单词
    - phonetic: 音标
    - translation: 中文翻译
    """
    if not request.ids:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="请选择要导出的生词"
        )

    # 获取生词数据
    stmt = select(Vocabulary).where(
        and_(
            Vocabulary.id.in_(request.ids),
            Vocabulary.user_id == current_user.id
        )
    )
    result = await db.execute(stmt)
    vocab_list = result.scalars().all()

    if not vocab_list:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="未找到要导出的生词"
        )

    # 创建Word文档
    doc = Document()

    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("生词本默写练习")
    title_run.bold = True
    title_run.font.size = Pt(22)  # 2号
    set_run_font(title_run, TITLE_FONT_NAME, Pt(22))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 添加表头
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    # 设置表头样式
    hidden_fields = request.hidden_fields or []

    # 确定列顺序
    # 如果同时默写 word 和 translation，顺序为：音标、英文、中文
    # 如果只默写 translation，顺序为：英文、音标、中文
    # 其他情况默认：英文、音标、中文
    if 'word' in hidden_fields and 'translation' in hidden_fields:
        # 同时默写英文和中文：音标、英文、中文
        header_cells = table.rows[0].cells
        header_cells[0].text = '音标'
        header_cells[1].text = '英文单词'
        header_cells[2].text = '中文翻译'
    elif 'translation' in hidden_fields:
        # 只默写中文：英文、音标、中文
        header_cells = table.rows[0].cells
        header_cells[0].text = '英文单词'
        header_cells[1].text = '音标'
        header_cells[2].text = '中文翻译'
    else:
        # 默认：英文、音标、中文
        header_cells = table.rows[0].cells
        header_cells[0].text = '英文单词'
        header_cells[1].text = '音标'
        header_cells[2].text = '中文翻译'

    # 设置表头样式（居中、四号）
    for cell in header_cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(14)  # 四号
            set_run_font(cell.paragraphs[0].runs[0], TABLE_FONT_NAME, Pt(14))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置垂直居中
        set_cell_vertical_center(cell)

    # 添加数据行
    for vocab in vocab_list:
        row = table.add_row()
        row_cells = row.cells

        if 'word' in hidden_fields and 'translation' in hidden_fields:
            # 同时默写：音标、英文、中文
            row_cells[0].text = vocab.phonetic or ''  # 音标
            row_cells[1].text = ''  # 英文留空
            row_cells[2].text = ''  # 中文留空
        elif 'translation' in hidden_fields:
            # 只默写中文：英文、音标、中文
            row_cells[0].text = ''  # 英文留空
            row_cells[1].text = vocab.phonetic or ''  # 音标
            row_cells[2].text = ''  # 中文留空
        elif 'word' in hidden_fields:
            # 只默写英文：音标、英文、中文
            row_cells[0].text = vocab.phonetic or ''  # 音标
            row_cells[1].text = ''  # 英文留空
            row_cells[2].text = vocab.translation or ''  # 中文
        else:
            # 默认全部显示
            row_cells[0].text = vocab.word or ''
            row_cells[1].text = vocab.phonetic or ''
            row_cells[2].text = vocab.translation or ''

        # 设置数据行样式（居中、四号字体）
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置字体大小和样式
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(14)  # 四号
                set_run_font(cell.paragraphs[0].runs[0], TABLE_FONT_NAME, Pt(14))
            # 设置垂直居中
            set_cell_vertical_center(cell)

    # 保存到内存
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # 返回文件流
    from datetime import datetime
    from urllib.parse import quote
    filename = f"vocabulary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    return StreamingResponse(
        iter([file_stream.getvalue()]),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quote(filename)}'
        }
    )
